from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "MAS_University_Academic_Management_System_Documentation_UPDATED.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    style_table(table)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_diagram_page(doc, title, image_name):
    doc.add_page_break()
    image_path = ROOT / image_name
    use_landscape = True
    if image_path.exists():
        with Image.open(image_path) as image:
            width, height = image.size
            use_landscape = width >= height

    section = doc.sections[-1]
    if use_landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        max_image_width = 9.8
        max_image_height = 6.25
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        max_image_width = 6.5
        max_image_height = 8.35
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    doc.add_heading(title, level=1)
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        with Image.open(image_path) as image:
            width_px, height_px = image.size
            height_if_max_width = max_image_width * height_px / width_px
        if height_if_max_width <= max_image_height:
            paragraph.add_run().add_picture(str(image_path), width=Inches(max_image_width))
        else:
            paragraph.add_run().add_picture(str(image_path), height=Inches(max_image_height))
        add_caption(doc, title)
    else:
        doc.add_paragraph(f"Diagram file missing: {image_name}")
    doc.add_page_break()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def build_doc():
    doc = Document()
    configure_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("University Academic Management System")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("MAS Final Project Documentation")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Course", "MAS - Modeling and Analysis of Information Systems"],
            ["Project", "University Academic Management System"],
            ["Implementation", "Java, JavaFX, JPA/Hibernate, Maven, H2 database"],
            ["Actors", "Student, Teacher, Administrator, SystemScheduler"],
            ["Prepared for", "Final project documentation and defence"],
        ],
    )

    doc.add_heading("1. Project Scope", level=1)
    doc.add_paragraph(
        "The project is a University Academic Management System focused on attendance management, "
        "weekly schedules, class meeting details, notifications, room booking, reporting, and academic "
        "structure management. The system is intentionally business-domain oriented and contains more "
        "than twelve meaningful business classes."
    )
    doc.add_paragraph(
        "The system supports three human roles: Student, Teacher, and Administrator. Every human user "
        "logs in before accessing personal functionality, and the visible navigation depends on the role "
        "of the logged-in person. The SystemScheduler is modeled separately as an automated actor that "
        "processes scheduled notification tasks."
    )
    doc.add_paragraph(
        "The scope intentionally does not include every possible university process. Grading, payments, "
        "recruitment, dormitory management, and diploma processing are outside the project scope. The "
        "project is focused on attendance, schedules, reports, notifications, room booking, and related "
        "academic administration."
    )
    add_table(
        doc,
        ["Area", "Business classes"],
        [
            ["Users", "Person, Student, Teacher, Administrator, PersonEmail"],
            ["Academic structure", "Field, Semester, StudentGroup, Subject"],
            ["Scheduling", "WeeklyScheduleEntry, ClassMeeting, MeetingTime"],
            ["Attendance", "Attendance"],
            ["Reporting", "AttendanceReport, ReportLine"],
            ["Room booking", "Room, MeetingSlot, RoomBooking"],
            ["Notifications", "Notification, EmailNotification, SystemNotification, ScheduledNotificationTask"],
        ],
    )

    doc.add_heading("Actors", level=2)
    add_table(
        doc,
        ["Actor", "Description"],
        [
            ["Student", "Views public schedule, personal schedule, attendance history, group details, comments, and notifications."],
            ["Teacher", "Views assigned meetings, creates meetings for qualified subjects, comments meetings, and marks attendance under restrictions."],
            ["Administrator", "Manages meetings, weekly schedules, users, rooms, bookings, notifications, reports, and exports."],
            ["SystemScheduler", "Automated actor that processes scheduled notification tasks and creates reminder, warning, and report-ready notifications."],
        ],
    )

    doc.add_heading("2. User Stories", level=1)
    add_table(
        doc,
        ["ID", "Actor", "User story"],
        [
            ["US-01", "Person", "As a person, I want to log into the system so that I can access functions according to my role."],
            ["US-02", "Person", "As a person, I want to view the public weekly schedule so that I can see current university class meetings."],
            ["US-03", "Person", "As a person, I want to manage email addresses and change password so that my account data remains current."],
            ["US-04", "Person", "As a person, I want to view notifications and personal meeting history so that I can track academic events involving me."],
            ["US-05", "Student", "As a student, I want to view my upcoming and past class meetings so that I know my schedule and history."],
            ["US-06", "Student", "As a student, I want to view my attendance history and attendance comments so that I can monitor participation."],
            ["US-07", "Student", "As a student, I want to view my group details and classmates so that I understand my academic group."],
            ["US-08", "Teacher", "As a teacher, I want to view class meetings assigned to me so that I can manage teaching activity."],
            ["US-09", "Teacher", "As a teacher, I want to create meetings only for subjects I am qualified to teach."],
            ["US-10", "Teacher", "As a teacher, I want to mark attendance only for my assigned meetings and only after the meeting starts."],
            ["US-11", "Teacher", "As a teacher, I want to add class meeting and individual attendance comments."],
            ["US-12", "Administrator", "As an administrator, I want to create, edit, and cancel class meetings without destroying history."],
            ["US-13", "Administrator", "As an administrator, I want to create fixed weekly schedules and generate concrete class meetings for a semester."],
            ["US-14", "Administrator", "As an administrator, I want to manage rooms and bookings so that room conflicts are prevented."],
            ["US-15", "Administrator", "As an administrator, I want to generate attendance reports with single or combined filters and export them to CSV."],
            ["US-16", "Administrator", "As an administrator, I want report conclusions to show attendance performance in a readable way."],
            ["US-17", "Administrator", "As an administrator, I want to manage notifications and demonstrate the StudentGroup to Student association in the GUI."],
            ["US-18", "SystemScheduler", "As the SystemScheduler, I want to process pending notification tasks and create reminders, warnings, and report-ready messages."],
        ],
    )

    doc.add_heading("3. Functional Requirements", level=1)
    add_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR-01", "The system shall require login before personal functionality is available."],
            ["FR-02", "The system shall display role-based navigation for Student, Teacher, and Administrator."],
            ["FR-03", "The system shall display a public weekly schedule for all logged-in persons."],
            ["FR-04", "The system shall restrict full class meeting details to involved students, assigned teachers, and administrators."],
            ["FR-05", "The system shall allow students to view personal schedule, attendance, group, notifications, and history."],
            ["FR-06", "The system shall allow teachers to view assigned meetings, create qualified meetings, comment meetings, and mark attendance."],
            ["FR-07", "The system shall allow administrators to manage meetings, weekly schedules, rooms, bookings, users, reports, and notifications."],
            ["FR-08", "The system shall enforce global uniqueness of email addresses."],
            ["FR-09", "The system shall allow defining subjects a teacher is qualified to teach, from one to five subjects."],
            ["FR-10", "The system shall prevent class meeting creation if the teacher is not qualified for the selected subject."],
            ["FR-11", "The system shall prevent attendance marking for cancelled class meetings."],
            ["FR-12", "The system shall allow attendance status and optional individual comments per student."],
            ["FR-13", "The system shall support flexible attendance report filters."],
            ["FR-14", "The system shall display report lines and an overall attendance performance conclusion."],
            ["FR-15", "The system shall export generated non-empty reports to CSV."],
            ["FR-16", "The system shall manage room bookings using MeetingSlot as a qualifier."],
            ["FR-17", "The system shall process scheduled notification tasks through SystemScheduler."],
            ["FR-18", "The system shall provide GroupStudentsAssociationView to demonstrate StudentGroup -> Students interaction."],
            ["FR-19", "The public weekly schedule shall show subject, class type, group, teacher, teacher email, date, time, mode, location or online link, and status."],
            ["FR-20", "A teacher shall be able to mark attendance only for assigned meetings after the meeting start time."],
            ["FR-21", "The system shall retrieve attendance students through ClassMeeting -> StudentGroup -> Students."],
            ["FR-22", "The system shall prevent duplicate attendance records for the same student and class meeting."],
            ["FR-23", "The system shall allow administrators to create weekly schedule templates and generate concrete class meetings."],
            ["FR-24", "The system shall display personal weekly schedule templates to students and assigned weekly entries to teachers."],
            ["FR-25", "The system shall allow administrators to create, edit, and cancel notifications, except sent notification content."],
            ["FR-26", "The system shall require classroom meetings to have room or location data and online meetings to have an online link before scheduling."],
        ],
    )

    doc.add_heading("4. Non-Functional Requirements", level=1)
    add_bullets(
        doc,
        [
            "The application shall be implemented in Java using JavaFX, JPA/Hibernate, and Maven.",
            "The application shall not use Spring.",
            "All persistent business data shall be stored using JPA/Hibernate.",
            "The GUI shall be understandable during defence and shall visibly show role-based access control.",
            "The GUI shall include sample data for Student, Teacher, and Administrator demo accounts.",
            "The application shall validate user input and show clear error messages.",
            "The implementation shall preserve full getters and setters in entity and embeddable classes.",
            "The implementation shall avoid obvious comments and keep business logic in services.",
            "Write operations shall be executed in transactions and rolled back on failure.",
            "Lazy associations shall be kept by default, with join fetch queries used where GUI screens require related data.",
            "The application shall be runnable using Maven and JavaFX.",
            "Diagrams in documentation shall be labeled, readable, and exported in lossless/vector form.",
        ],
    )

    add_diagram_page(doc, "5. Use Case Diagram", "USE_CASE_DIAGRAM.png")

    doc.add_heading("6. Non-Trivial Use Case Scenarios", level=1)
    doc.add_heading("6.1. Generate and export attendance report", level=2)
    doc.add_paragraph(
        "This use case is non-trivial because it references report generation, attendance aggregation, "
        "optional filters, report persistence, conclusion generation, and CSV export. Export attendance "
        "report extends Generate attendance report."
    )
    add_numbers(
        doc,
        [
            "Administrator logs into the system.",
            "Administrator opens the Generate Reports screen.",
            "System displays optional filters: semesters, teacher, subject, group, class type, date from, and date to.",
            "Administrator selects any combination of filters or leaves filters empty.",
            "Administrator clicks Generate Report.",
            "System finds all class meetings matching the filter.",
            "System finds students through ClassMeeting -> StudentGroup -> Students.",
            "System loads attendance records for matching meetings.",
            "System calculates total meetings, present, late, excused, absent, and attendance percentage for each student.",
            "System creates AttendanceReport and ReportLine objects.",
            "System calculates overallPerformancePercentage as the average of report line percentages.",
            "System displays the report table and conclusion message.",
            "Administrator clicks Export CSV.",
            "System validates that a report exists and contains lines.",
            "System writes report metadata, conclusion, and report lines to a CSV file.",
        ],
    )
    doc.add_heading("Alternative flows", level=2)
    add_bullets(
        doc,
        [
            "If no report has been generated, the system displays: Generate a report before exporting.",
            "If the generated report has no lines, the system displays: Cannot export an empty report.",
            "If a filter is null or empty, the system treats it as unrestricted.",
        ],
    )

    doc.add_heading("6.2. Mark attendance", level=2)
    doc.add_paragraph(
        "This use case is also non-trivial because it references opening class meeting details, "
        "role-based authorization, class meeting status validation, meeting start time validation, "
        "association traversal through ClassMeeting -> StudentGroup -> Students, duplicate attendance "
        "prevention, and persistent attendance comments."
    )
    add_numbers(
        doc,
        [
            "Teacher or Administrator logs into the system.",
            "User opens Public Weekly Schedule or a personal class meeting list.",
            "User opens a class meeting details window.",
            "System verifies that the user is authorized to open the full details.",
            "User clicks Mark Attendance.",
            "System checks that the class meeting is not cancelled.",
            "If the user is a teacher, the system checks that the teacher is assigned to the meeting.",
            "If the user is a teacher, the system checks that the meeting start time has already passed.",
            "System loads students through ClassMeeting -> StudentGroup -> Students.",
            "System displays one attendance row per student.",
            "User selects an attendance status for each student and optionally enters an individual comment.",
            "System validates that every row has a status.",
            "System checks that duplicate attendance records do not already exist for the same student and meeting.",
            "System saves attendance records in a transaction.",
            "System displays a success message and the attendance becomes visible in student history and reports.",
        ],
    )
    doc.add_heading("Alternative flows for mark attendance", level=2)
    add_bullets(
        doc,
        [
            "If the teacher is not assigned to the meeting, the system displays: You are not assigned to this class meeting.",
            "If the meeting has not started yet, the system displays: Attendance cannot be marked before the class meeting starts.",
            "If the meeting is cancelled, the system displays: Cannot mark attendance for a cancelled class meeting.",
            "If a status is missing, the system displays: Attendance status is required.",
            "If duplicate attendance exists, the system displays: Attendance has already been registered for this student and meeting.",
        ],
    )

    add_diagram_page(doc, "7. Activity Diagram - Generate And Export Attendance Report", "ACTIVITY_GENERATE_AND_EXPORT_REPORT.png")
    add_diagram_page(doc, "8. Activity Diagram - Mark Attendance", "ACTIVITY_MARK_ATTENDANCE.png")
    add_diagram_page(doc, "9. State Diagram - ClassMeeting Lifecycle", "STATE_CLASS_MEETING.png")
    add_diagram_page(doc, "10. Analytical Class Diagram", "ANALYTICAL_CLASS_DIAGRAM_COMPACT.png")
    add_diagram_page(doc, "11. Design Class Diagram", "DESIGN_CLASS_DIAGRAM_COMPACT.png")

    doc.add_heading("12. GUI Design", level=1)
    doc.add_paragraph(
        "The selected GUI design for the non-trivial report use case is the Administrator Generate Reports screen. "
        "The screen contains filter controls, a report table, a conclusion label, and an export button."
    )
    doc.add_heading("Role-based main layout", level=2)
    add_table(
        doc,
        ["Role", "Visible navigation"],
        [
            ["All logged-in persons", "Public Weekly Schedule, Personal Settings, My Notifications, My History, Logout"],
            ["Student", "My Schedule, My Attendance, My Group"],
            ["Teacher", "My Class Meetings, Create Class Meeting, Group Students Association"],
            ["Administrator", "Manage Class Meetings, Manage Weekly Schedules, Manage Users, Manage Notifications, Generate Reports, Export Reports, Manage Rooms, Group Students Association"],
        ],
    )
    doc.add_heading("Administrator report screen", level=2)
    add_table(
        doc,
        ["GUI element", "Purpose"],
        [
            ["Semesters multi-select", "Allows filtering by one or more semesters."],
            ["Teacher combo box", "Filters class meetings by teacher."],
            ["Subject combo box", "Filters class meetings by subject."],
            ["Group combo box", "Filters by student group."],
            ["Class type combo box", "All, Lecture, Tutorial, or Laboratory."],
            ["Date fields", "Optional date range filter."],
            ["Generate Report button", "Creates and displays the report."],
            ["Report table", "Shows student, total meetings, present, late, excused, absent, and attendance percentage."],
            ["Conclusion label", "Shows the overall attendance performance message."],
            ["Export CSV button", "Exports a generated non-empty report."],
        ],
    )
    doc.add_heading("Class meeting details and attendance marking", level=2)
    doc.add_paragraph(
        "The ClassMeetingDetailsView displays subject, class type, group, teacher, teacher email, date, time, "
        "meeting mode, room or online link, status, and general comment. Student access is read-only. Teachers "
        "may edit comments and mark attendance only for assigned meetings when validation rules pass. "
        "Administrators may edit, cancel, comment, and mark attendance for any meeting."
    )
    add_table(
        doc,
        ["Attendance table column", "Meaning"],
        [
            ["Student number", "Identifier of the student from the selected meeting group."],
            ["Student name", "Student retrieved through ClassMeeting -> StudentGroup -> Students."],
            ["Attendance status", "PRESENT, ABSENT, LATE, or EXCUSED."],
            ["Individual comment", "Optional editable comment visible in attendance history."],
        ],
    )
    doc.add_heading("GUI requirement 4.2.4", level=2)
    doc.add_paragraph(
        "The implementation also contains GroupStudentsAssociationView. It displays StudentGroup objects in a ListView "
        "and displays multiple Student objects in a TableView after a group is selected. Students are retrieved through "
        "StudentGroup.getStudents(), not by filtering all students or using a query on selection change."
    )

    doc.add_heading("13. Dynamic Analysis Implications", level=1)
    add_table(
        doc,
        ["Dynamic analysis result", "Implementation/design consequence"],
        [
            ["Cancelled meetings must remain visible", "ClassMeetingStatus includes CANCELLED; meetings are not deleted."],
            ["Draft meetings may be incomplete", "ClassMeetingStatus includes DRAFT."],
            ["Teacher cannot mark attendance early", "ClassMeetingService.validateCanMarkAttendance checks meeting start time."],
            ["Missing report filters mean unrestricted criteria", "AttendanceReportFilter uses nullable fields."],
            ["Reports need a conclusion", "AttendanceReport stores overallPerformancePercentage and generates a conclusion message."],
            ["Room booking conflicts depend on exact slot", "MeetingSlot is an embeddable qualifier and RoomBooking has a unique room-slot constraint."],
            ["Notifications require a business trigger", "ScheduledNotificationTask is processed by SystemScheduler."],
        ],
    )

    doc.add_heading("14. Design And Implementation Decisions", level=1)
    add_table(
        doc,
        ["Decision", "Explanation"],
        [
            ["Persistence", "JPA/Hibernate with RESOURCE_LOCAL transactions and H2 database are used for persistent data."],
            ["Inheritance", "Person is implemented with JOINED inheritance for Student, Teacher, and Administrator."],
            ["Notification inheritance", "Notification uses SINGLE_TABLE inheritance for EmailNotification and SystemNotification."],
            ["Multi-value emails", "Person.emails is implemented as an ElementCollection in person_emails."],
            ["Teacher qualification", "Teacher and Subject use many-to-many teacher_subject association."],
            ["Association class", "Attendance connects Student and ClassMeeting while storing status and comment."],
            ["Qualified association", "Room uses MeetingSlot as qualifier for bookingsBySlot."],
            ["Design class diagram", "Unsupported analytical constructs are replaced by Java classes, embeddables, services, and repositories."],
            ["Manual architecture", "The project uses JavaFX, services, repositories, and JPA manually; Spring is not used."],
        ],
    )

    doc.add_heading("15. Implementation Structure", level=1)
    doc.add_paragraph(
        "The application is implemented in Java. JavaFX provides the graphical user interface, "
        "JPA/Hibernate provides persistence, and Maven manages dependencies and execution. The project "
        "uses manual services and repositories rather than Spring."
    )
    add_table(
        doc,
        ["Package", "Responsibility"],
        [
            ["model", "Persistent entities, embeddable value objects, and enums."],
            ["persistence", "Manual repository classes and JpaUtil for EntityManagerFactory access."],
            ["service", "Business logic, validation, authentication, scheduling, reporting, exporting, and notifications."],
            ["app", "JavaFX application, screens, table row view models, and role-based navigation."],
        ],
    )
    doc.add_paragraph(
        "The application starts with the login screen. After successful authentication, the main layout "
        "displays navigation according to the logged-in role. Important business workflows, such as "
        "attendance registration and report generation, are implemented through service classes that "
        "coordinate repositories and validation rules."
    )

    doc.add_heading("16. Implementation Compliance", level=1)
    add_table(
        doc,
        ["Requirement", "How the implementation satisfies it"],
        [
            ["4.2.1 complete class structure", "Entities include users, academic structure, scheduling, attendance, reports, rooms, and notifications."],
            ["4.2.2 selected use case methods", "Report generation/export, attendance marking, schedule generation, and access validation are implemented in services."],
            ["4.2.3 GUI", "JavaFX login, role-based navigation, schedules, reports, rooms, notifications, and settings are implemented."],
            ["4.2.4 association GUI", "GroupStudentsAssociationView uses ListView and TableView over StudentGroup.getStudents()."],
            ["4.2.5 sample data", "SampleDataService creates demo Student, Teacher, Administrator, group, subject, room, schedules, meetings, notifications."],
            ["4.2.6 GUI usability", "Role-specific navigation limits visible functions and provides clear messages."],
            ["4.2.7 persistence", "H2 database with Hibernate persists all entities."],
            ["4.2.8 comments", "The code avoids unnecessary obvious comments."],
            ["4.2.10 technology", "The project is implemented in Java."],
        ],
    )

    doc.add_heading("17. Demo Accounts", level=1)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ["Student", "anna.nowak@student.pja.edu.pl", "student"],
            ["Teacher", "piotr.kowalski@pja.edu.pl", "teacher"],
            ["Administrator", "admin@pja.edu.pl", "admin"],
        ],
    )

    doc.add_heading("18. Defence Demonstration Scenario", level=1)
    doc.add_heading("Student demonstration", level=2)
    add_numbers(
        doc,
        [
            "Log in as Student.",
            "Show the public weekly schedule.",
            "Try opening a meeting not involving the student and show the public-only access message.",
            "Open My Schedule and show upcoming and past class meetings.",
            "Open My Attendance and show attendance history and comments.",
            "Open My Group and show group details, classmates, and group schedule.",
            "Open Personal Settings and My Notifications.",
        ],
    )
    doc.add_heading("Teacher demonstration", level=2)
    add_numbers(
        doc,
        [
            "Log in as Teacher.",
            "Show teacher-specific navigation.",
            "Open assigned class meetings and class meeting details.",
            "Add or edit a general class meeting comment.",
            "Try marking attendance for a future meeting and show that it is blocked.",
            "Open a current or past assigned meeting and mark attendance.",
            "Add individual attendance comments and save attendance.",
            "Explain that students are loaded through ClassMeeting -> StudentGroup -> Students.",
        ],
    )
    doc.add_heading("Administrator demonstration", level=2)
    add_numbers(
        doc,
        [
            "Log in as Administrator.",
            "Show admin navigation and unrestricted class meeting detail access.",
            "Create or edit a weekly schedule entry and generate class meetings.",
            "Cancel a meeting and explain that it is not deleted.",
            "Generate an attendance report using combined filters.",
            "Show report lines and the attendance performance conclusion.",
            "Export the report to CSV.",
            "Open room management and explain Room [MeetingSlot] -> RoomBooking.",
            "Open notification management and explain SystemScheduler tasks.",
            "Open GroupStudentsAssociationView to demonstrate MAS requirement 4.2.4.",
        ],
    )

    doc.add_heading("19. Requirement Coverage Checklist", level=1)
    add_table(
        doc,
        ["Requirement", "Included"],
        [
            ["User stories", "Yes - Section 2"],
            ["Use case diagram", "Yes - Section 5"],
            ["Non-functional requirements", "Yes - Section 4"],
            ["Analytical class diagram", "Yes - Section 9"],
            ["Design class diagram", "Yes - Section 10"],
            ["Use case scenario enumeration", "Yes - Section 6"],
            ["Activity diagram", "Yes - Sections 7 and 8"],
            ["State diagram", "Yes - Section 9"],
            ["GUI design", "Yes - Section 12"],
            ["Design decisions and dynamic analysis discussion", "Yes - Sections 13 and 14"],
            ["Implementation structure", "Yes - Section 15"],
            ["Implementation requirement 4.2.4", "Yes - GroupStudentsAssociationView"],
            ["Defence scenario", "Yes - Section 18"],
        ],
    )

    doc.add_heading("20. Known Limitations", level=1)
    add_bullets(
        doc,
        [
            "The SharePoint version of the document was not accessible without Microsoft login during this revision.",
            "The generated diagrams should be reviewed visually after insertion into the final PDF to ensure they remain readable at 100% zoom.",
            "If the institution requires a specific filename, export the final PDF as MAS_Group_Lastname_Firstname_StudentNo.pdf.",
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
