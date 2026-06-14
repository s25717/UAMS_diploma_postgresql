from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "DIPLOMA_DOCUMENTATION.md"
OUT = ROOT / "docs" / "UAMS_Diploma_Documentation_s25717_Zhanel_Nurkhabanova.docx"

FONT = "Times New Roman"
BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(90, 90, 90)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None,
                 color: RGBColor | None = None, font: str = FONT) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    p._p.append(field)


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field to refresh the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 1", 14, 18, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.footer_distance = Cm(1.25)


def add_centered(doc: Document, text: str, size: float = 12, bold: bool = False,
                 italic: bool = False, after: float = 6, color: RGBColor = BLACK) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def add_title_page(doc: Document) -> None:
    add_centered(doc, "Polish-Japanese Academy of Information Technology", 14, bold=True, after=3)
    add_centered(doc, "Faculty: [to be completed]", 12, after=3)
    add_centered(doc, "Department: [to be completed]", 12, after=3)
    add_centered(doc, "Specialization: Databases and Database Design", 12, after=48)

    add_centered(doc, "Zhanel Nurkhabanova", 14, bold=True, after=3)
    add_centered(doc, "Student ID: s25717", 12, after=42)

    add_centered(doc, "University Academic Management System (UAMS)", 18, bold=True, after=8)
    add_centered(doc, "PostgreSQL Database Architecture and Application Prototype", 14, bold=True, after=42)

    add_centered(doc, "Diploma thesis / diploma project", 12, after=10)
    add_centered(doc, "Supervisor: [to be completed]", 12, after=72)
    add_centered(doc, "Warsaw, June 2026", 12, after=0)
    doc.add_page_break()


def add_polish_title_page(doc: Document) -> None:
    add_centered(doc, "Polsko-Japońska Akademia Technik Komputerowych", 14, bold=True, after=3)
    add_centered(doc, "Nazwa Wydziału: [uzupełnić]", 12, after=3)
    add_centered(doc, "Nazwa Katedry: [uzupełnić]", 12, after=3)
    add_centered(doc, "Specjalizacja: Bazy danych i projektowanie baz danych", 12, after=48)

    add_centered(doc, "Zhanel Nurkhabanova", 14, bold=True, after=3)
    add_centered(doc, "Nr albumu: s25717", 12, after=42)

    add_centered(doc, "UAMS: Uniwersytecki system zarządzania akademickiego", 18, bold=True, after=8)
    add_centered(doc, "Architektura bazy danych PostgreSQL i prototyp aplikacji", 14, bold=True, after=42)

    add_centered(doc, "Rodzaj pracy dyplomowej: [uzupełnić]", 12, after=10)
    add_centered(doc, "Promotor: [uzupełnić]", 12, after=72)
    add_centered(doc, "Warszawa, czerwiec 2026", 12, after=0)


def add_abstracts(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    add_body(doc, (
        "This thesis documentation describes the University Academic Management System, a JavaFX desktop prototype "
        "supported by a PostgreSQL database. The project focuses on explicit relational database design, Flyway "
        "migrations, database constraints, triggers, indexes, and proof scripts for important business rules. "
        "The system covers users, academic structure, schedules, room bookings, attendance, reports, notifications, "
        "and activity history. Special attention is given to semester-field curriculum modelling and to enforcing "
        "data correctness directly in PostgreSQL."
    ))
    add_body(doc, "Keywords: PostgreSQL, database design, Flyway, JavaFX, academic management system")

    doc.add_heading("Streszczenie", level=1)
    add_body(doc, (
        "Dokumentacja opisuje system University Academic Management System, czyli prototyp aplikacji desktopowej "
        "JavaFX wspierany przez bazę danych PostgreSQL. Projekt koncentruje się na jawnie zaprojektowanej "
        "relacyjnej bazie danych, migracjach Flyway, ograniczeniach, wyzwalaczach, indeksach oraz skryptach "
        "potwierdzających działanie reguł biznesowych. System obejmuje użytkowników, strukturę akademicką, plany "
        "zajęć, rezerwacje sal, obecności, raporty, powiadomienia oraz historię aktywności."
    ))
    add_body(doc, "Słowa kluczowe: PostgreSQL, projektowanie baz danych, Flyway, JavaFX, system akademicki")
    doc.add_page_break()


def add_toc_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    p.add_run("Table of Contents")
    set_run_font(p.runs[0], size=14, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    add_toc(p)
    doc.add_page_break()


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=10.5, font="Courier New")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=12)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    set_run_font(run, size=9.5, font="Courier New")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, italic=True)


def add_markdown_table(doc: Document, rows: list[list[str]], caption: str | None) -> None:
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    col_count = len(rows[0])
    if col_count == 2:
        widths = [2600, 6760]
    elif col_count == 3:
        widths = [2200, 3580, 3580]
    else:
        widths = [int(9360 / col_count)] * col_count
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F2F2")
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(table_lines):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if idx == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def table_caption(rows: list[list[str]], chapter: int, counts: defaultdict[int, int]) -> str | None:
    if not rows:
        return None
    header = [h.lower() for h in rows[0]]
    counts[chapter] += 1
    number = f"{chapter}.{counts[chapter]}"
    if header == ["decision", "pros", "cons"]:
        return f"Table {number}: Summary of architectural choices."
    if header == ["area", "technology"]:
        return f"Table {number}: Technology stack used in the system."
    return f"Table {number}: Structured project information."


def convert_markdown_body(doc: Document) -> None:
    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    lines = raw_lines[2:] if raw_lines and raw_lines[0].startswith("# ") else raw_lines

    h1 = 0
    h2 = 0
    h3 = 0
    table_counts: defaultdict[int, int] = defaultdict(int)
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("## "):
            h1 += 1
            h2 = 0
            h3 = 0
            doc.add_heading(f"{h1}. {line[3:].strip()}", level=1)
            i += 1
            continue
        if line.startswith("### "):
            h2 += 1
            h3 = 0
            doc.add_heading(f"{h1}.{h2}. {line[4:].strip()}", level=2)
            i += 1
            continue
        if line.startswith("#### "):
            h3 += 1
            doc.add_heading(f"{h1}.{h2}.{h3}. {line[5:].strip()}", level=3)
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows, table_caption(rows, max(h1, 1), table_counts))
            continue

        stripped = line.strip()
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        add_body(doc, line)
        i += 1


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = [
        '[1] "Niezbędnik pracy dyplomowej," Polish-Japanese Academy of Information Technology. '
        "[Online]. Available: https://pja.edu.pl/niezbednik-pracy-dyplomowej-wzi/. [Accessed: 14.06.2026].",
        '[2] "UAMS Diploma thesis," Notion. [Online]. Available: '
        "https://app.notion.com/p/UAMS-Diploma-thesis-3754ebf0d4738021bf1afd810d004c58. [Accessed: 14.06.2026].",
        '[3] "UAMS diploma PostgreSQL repository," GitHub repository s25717/UAMS_diploma_postgresql. '
        "[Online]. Available: https://github.com/s25717/UAMS_diploma_postgresql. [Accessed: 14.06.2026].",
    ]
    for ref in refs:
        add_body(doc, ref)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    set_page_setup(doc)

    add_title_page(doc)
    add_polish_title_page(doc)
    content_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_setup(doc)
    content_section.footer.is_linked_to_previous = False
    restart_page_numbering(content_section, 1)
    add_page_number(content_section)
    add_abstracts(doc)
    add_toc_page(doc)
    convert_markdown_body(doc)
    add_references(doc)
    enable_update_fields(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
